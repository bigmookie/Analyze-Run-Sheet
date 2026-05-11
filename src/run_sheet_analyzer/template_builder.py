"""Programmatically prepare the firm's report template for docxtpl.

Reads `Abstract - Report - Minerals.docx`, performs targeted text substitutions
to introduce Jinja tags for the certificate-page fields, and deletes the
existing single-vesting TITLE REPORT block (the renderer builds per-tract
sections programmatically after docxtpl renders the certificate). Saves the
result to `templates/title-report.docx`.
"""
from __future__ import annotations

import hashlib
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentT


# Literal source-template tokens → Jinja replacement tags. Applied to each
# paragraph's full text via the simple-replacement helper, so we never have to
# touch run boundaries directly.
INLINE_REPLACEMENTS: list[tuple[str, str]] = [
    ("ADDRESSEE", "{{ addressee }}"),
    ("Effective Date: EFFECTIVE DATE", "Effective Date: {{ effective_date }}"),
    ("Records of COUNTY County", "Records of {{ county }} County"),
    # The signing-date line in the firm template embeds a literal date; replace it.
    # We match flexibly on any "on this the <date>," form.
]

# Paragraph after which we delete everything (inclusive). The firm template has
# a single TITLE REPORT block we don't want — we generate per-tract sections in
# Python after docxtpl finishes the certificate page.
DELETE_FROM_HEADING = "TITLE REPORT"


def _replace_in_paragraph(paragraph, old: str, new: str) -> bool:
    """Replace `old` with `new` in a paragraph's text, preserving the first run's
    formatting. Returns True if a replacement happened."""
    text = paragraph.text
    if old not in text:
        return False
    new_text = text.replace(old, new)
    # Collapse all runs into the first, set the new text. This loses fine-grained
    # formatting across runs but for the certificate fields that's acceptable.
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return True
    first = paragraph.runs[0]
    first.text = new_text
    for r in paragraph.runs[1:]:
        r.text = ""
    return True


def _replace_signing_date(paragraph) -> bool:
    """The certificate has 'Signed at Madison, Mississippi, on this the <date>,'.
    Replace the date span with `{{ signing_date }}`."""
    import re
    text = paragraph.text
    m = re.search(r"on this the\s+(.+?),\s*effective", text)
    if not m:
        return False
    new_text = text[: m.start(1)] + "{{ signing_date }}" + text[m.end(1):]
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for r in paragraph.runs[1:]:
            r.text = ""
    else:
        paragraph.add_run(new_text)
    return True


def _delete_from_heading(doc: DocumentT, heading: str) -> None:
    body = doc.element.body
    deleting = False
    to_remove = []
    for child in list(body.iterchildren()):
        # Preserve the document's terminal section-properties element regardless.
        if child.tag.endswith("}sectPr"):
            continue
        if not deleting:
            text = "".join(t.text or "" for t in child.iter() if t.tag.endswith("}t"))
            if text.strip() == heading:
                deleting = True
        if deleting:
            to_remove.append(child)
    for child in to_remove:
        body.remove(child)


def source_hash(source: Path) -> str:
    return hashlib.sha256(source.read_bytes()).hexdigest()


def build_template(source: Path, dest: Path) -> None:
    doc = Document(source)
    for paragraph in doc.paragraphs:
        for old, new in INLINE_REPLACEMENTS:
            _replace_in_paragraph(paragraph, old, new)
        _replace_signing_date(paragraph)
    _delete_from_heading(doc, DELETE_FROM_HEADING)
    dest.parent.mkdir(parents=True, exist_ok=True)
    doc.save(dest)
    # Stash the source hash next to the dest so we can detect changes later.
    (dest.parent / (dest.stem + ".source.sha256")).write_text(source_hash(source))


def needs_rebuild(source: Path, dest: Path) -> bool:
    if not dest.exists():
        return True
    sidecar = dest.parent / (dest.stem + ".source.sha256")
    if not sidecar.exists():
        return True
    try:
        return sidecar.read_text().strip() != source_hash(source)
    except Exception:
        return True


def ensure_template(source: Path, dest: Path) -> Path:
    if needs_rebuild(source, dest):
        build_template(source, dest)
    return dest
