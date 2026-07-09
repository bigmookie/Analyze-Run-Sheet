"""Render the consolidated draft report.

  1. docxtpl fills the certificate page from templates/title-report.docx.
  2. python-docx appends each tract's report-section text (as produced by
     the analyzer) one paragraph at a time, with light formatting.
"""
from __future__ import annotations

import io
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docxtpl import DocxTemplate

from .analyzer import JobConfig


CALLOUT_RED = RGBColor(0xC0, 0x00, 0x00)

# Shown on the certificate when a job field is left blank, so the examiner can
# find-and-fill them in the generated .docx.
_PLACEHOLDERS = {
    "addressee": "[ADDRESSEE]",
    "effective_date": "[EFFECTIVE DATE]",
    "county": "[COUNTY]",
    "signing_date": "[SIGNING DATE]",
}


def _add_heading_run(doc, text: str, size: int = 14):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    return p


def _add_paragraph(
    doc, text: str, *,
    bold: bool = False, italic: bool = False, color=None, size: int | None = None,
    space_before: int = 0, space_after: int = 4,
    left_indent: float | None = None, hanging: float | None = None,
):
    """Add one formatted paragraph. Spacing is in points; indents in inches.
    Whitespace is controlled by paragraph spacing, not blank paragraphs."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if color is not None:
        r.font.color.rgb = color
    if size is not None:
        r.font.size = Pt(size)
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if left_indent is not None:
        pf.left_indent = Inches(left_indent)
    if hanging is not None:
        pf.first_line_indent = Inches(-hanging)
    return p


# Top-level report sections, rendered as bold headings with space above.
_SECTION_HEADS = {"VESTING", "DESCRIPTION", "EXCEPTIONS", "ATTORNEY REVIEW"}
# VESTING sub-labels, rendered bold (accepted with or without a trailing colon).
_SUB_LABELS = {"Surface", "Minerals"}


def _render_tract_section(doc, tract_id: str, text: str, job: JobConfig):
    """Render one tract's report section into the doc.

    The model produces a plain-text section in the firm's house style. We format
    it for readability:
      - Section headings (VESTING, DESCRIPTION, EXCEPTIONS, ATTORNEY REVIEW):
        bold, larger, with space above.
      - Sub-labels (Surface, Minerals): bold.
      - Numbered subsection headers (e.g. "1. Voluntary Liens (...):"): bold.
      - Bullet lines ("- " / "• "): real bullets with a hanging indent.
      - The minerals "Total ..." reconciliation line: italic, indented.
      - Truncation callouts ("*** "): red bold.
    Blank input lines are dropped; whitespace comes from paragraph spacing.
    """
    doc.add_page_break()
    _add_heading_run(doc, f"TITLE REPORT — Tract {tract_id}", size=14)

    for raw_line in text.splitlines():
        stripped = raw_line.strip()

        # Skip blanks (spacing is handled per-paragraph) and the model's header.
        if not stripped:
            continue
        if stripped.startswith("=== TRACT") and stripped.endswith("==="):
            continue

        # Defensively drop County Taxes — removed from the report. Guards against
        # stale cached output or model drift that still emits it.
        if stripped.endswith("County Taxes:") or "[See parcel data" in stripped:
            continue

        # Section headings.
        if stripped in _SECTION_HEADS:
            _add_paragraph(doc, stripped, bold=True, size=13, space_before=12, space_after=4)
            continue

        # VESTING sub-labels (Surface / Minerals), with or without a colon.
        if stripped.rstrip(":") in _SUB_LABELS:
            _add_paragraph(doc, stripped.rstrip(":"), bold=True, space_before=6, space_after=2)
            continue

        # Numbered subsection headers (e.g. "1. Voluntary Liens (...):").
        if stripped[:2].rstrip(".").isdigit() and stripped.endswith(":"):
            _add_paragraph(doc, stripped, bold=True, space_before=8, space_after=2)
            continue

        # Truncation callout.
        if stripped.startswith("*** "):
            _add_paragraph(doc, stripped, bold=True, color=CALLOUT_RED, space_before=4)
            continue

        # Bullet lines → real bullet with hanging indent.
        if stripped.startswith("- ") or stripped.startswith("• "):
            item = stripped[2:].strip()
            _add_paragraph(doc, f"• {item}", left_indent=0.35, hanging=0.2, space_after=3)
            continue

        # Minerals reconciliation line.
        if stripped.startswith("Total "):
            _add_paragraph(doc, stripped, italic=True, left_indent=0.2, space_before=2, space_after=4)
            continue

        # Everything else (e.g. DESCRIPTION paragraphs).
        _add_paragraph(doc, stripped, space_after=4)


def render_report(
    *,
    template_path: Path,
    output_path: Path,
    job: JobConfig,
    tract_sections: list[tuple[str, str]],   # [(tract_id, report_text), ...]
) -> Path:
    # Build the whole document in memory, then write to disk exactly once.
    # Writing the destination twice (tpl.save → reopen → doc.save) is fragile
    # in cloud-synced folders (Dropbox/OneDrive): the sync client memory-maps
    # the file created by the first save to hash/upload it, and the second
    # save's truncating reopen then fails with OSError errno 22 (Windows
    # ERROR_USER_MAPPED_FILE). A single create-write closes that window.
    buffer = io.BytesIO()

    # Certificate page from the prepared template.
    tpl = DocxTemplate(str(template_path))
    tpl.render({
        "addressee": job.addressee or _PLACEHOLDERS["addressee"],
        "effective_date": job.effective_date or _PLACEHOLDERS["effective_date"],
        "county": job.county or _PLACEHOLDERS["county"],
        "signing_date": job.signing_date or _PLACEHOLDERS["signing_date"],
    })
    tpl.save(buffer)

    # Append per-tract sections.
    buffer.seek(0)
    doc = Document(buffer)
    for tract_id, text in tract_sections:
        _render_tract_section(doc, tract_id, text, job)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
